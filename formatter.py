from docx import Document
from fpdf import FPDF
import io


def generate_docx(sections):
    doc = Document()
    doc.add_heading("AI-Assisted Academic Document", level=0)

    for title, content in sections.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(sections):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(0, 10, "AI-Assisted Academic Document", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", size=11)

    for title, content in sections.items():
        pdf.ln(5)
        pdf.set_font("Arial", style="B", size=12)
        pdf.multi_cell(0, 8, title)

        pdf.ln(2)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 8, content)

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer
